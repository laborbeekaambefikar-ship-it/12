"""
Convert Dissertation_Arbaz.pdf -> Dissertation_Arbaz.docx
applying the dissertation formatting guidelines:

  - Times New Roman, A4 paper with 1-inch margins
  - Body text: 12 pt, 1.5 line spacing, justified, 12 pt after
  - Main headings: 14 pt bold, left aligned, 12 pt before & after
  - Sub-headings: 12 pt bold, left aligned
  - Figure / Table captions: 12 pt bold, centered, 12 pt above & below

Strategy: run pdf2docx first to get text + tables + images in reading
order, then walk the produced document and rewrite styling per spec.
"""

import re
import shutil
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PDF_FILE = '/projects/sandbox/12/Dissertation_Arbaz.pdf'
DOCX_FILE = '/projects/sandbox/12/Dissertation_Arbaz.docx'

# ---------------------------------------------------------------------------
# Step 1: fresh pdf2docx conversion
# ---------------------------------------------------------------------------
print('[1/3] Converting PDF to DOCX (pdf2docx)...')
cv = Converter(PDF_FILE)
cv.convert(DOCX_FILE, start=0, end=None)
cv.close()


# ---------------------------------------------------------------------------
# Step 2: helpers for restyling
# ---------------------------------------------------------------------------
W = qn


def set_run_font(run, name='Times New Roman', size_pt=None, bold=None):
    """Force every relevant font slot so Word actually uses TNR."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(W('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(W('w:ascii'), name)
    rFonts.set(W('w:hAnsi'), name)
    rFonts.set(W('w:eastAsia'), name)
    rFonts.set(W('w:cs'), name)
    run.font.name = name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
        # complex-script size mirror so non-Latin glyphs don't fall back
        szCs = rPr.find(W('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(W('w:val'), str(int(size_pt * 2)))
    if bold is not None:
        run.bold = bool(bold)


def get_max_size(paragraph):
    sizes = [r.font.size.pt for r in paragraph.runs if r.font.size]
    return max(sizes) if sizes else None


def all_runs_bold(paragraph):
    bs = [r.bold for r in paragraph.runs if r.text.strip()]
    return bool(bs) and all(bool(b) for b in bs)


CAPTION_RE = re.compile(r'^\s*(Figure|Fig\.?|Table)\s+\d', re.IGNORECASE)
# leader-dot patterns can be ".....", ". . . . .", or "...." with various spacing
TOC_DOTS_RE = re.compile(r'(\.\s*){3,}\.')
TRAILING_PAGE_RE = re.compile(r'(\.\s*){2,}\d+\s*$')


def is_toc_entry(text):
    """List of Figures / Tables entries: 'Figure 5.1 Some title ........ 23'."""
    t = text.strip()
    return bool(TOC_DOTS_RE.search(t) or TRAILING_PAGE_RE.search(t))


def is_caption(text):
    t = text.strip()
    if is_toc_entry(t):
        return False
    return bool(CAPTION_RE.match(t))


# Front-matter / chapter-level headings that the source PDF rendered
# at sizes < 16pt for layout reasons but are conceptually chapter titles.
KNOWN_MAIN_HEADINGS = {
    'declaration', 'certificate', 'acknowledgement', 'acknowledgements',
    'abstract', 'contents', 'table of contents',
    'list of tables', 'list of figures',
    'notations, symbols, and abbreviations',
    'bibliography', 'references', 'appendix',
    'complete state machine transition table',
    'software installation and launch guide',
}


# ---------- styling primitives ----------

def style_cover(p):
    """Cover-page items: keep centered, TNR, sized close to original."""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    src_max = get_max_size(p) or 12
    if src_max >= 17:
        sz = 16
    elif src_max >= 14:
        sz = 14
    else:
        sz = 12
    for r in p.runs:
        set_run_font(r, size_pt=sz, bold=(True if r.bold else r.bold))


def style_main_heading(p):
    """14 pt bold, left aligned, 12 pt before and after."""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5
    pf.keep_with_next = True
    for r in p.runs:
        set_run_font(r, size_pt=14, bold=True)


def style_sub_heading(p):
    """12 pt bold, left aligned."""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.keep_with_next = True
    for r in p.runs:
        set_run_font(r, size_pt=12, bold=True)


def style_caption(p):
    """12 pt bold, centered, 12 pt before & after."""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5
    for r in p.runs:
        set_run_font(r, size_pt=12, bold=True)


def style_body(p):
    """12 pt, TNR, 1.5 line, justified, 12 pt after."""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5
    for r in p.runs:
        set_run_font(r, size_pt=12, bold=r.bold)


def style_table_cell_paragraph(p):
    """Inside tables: TNR, 11 pt, tighter spacing for readability."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15
    for r in p.runs:
        set_run_font(r, size_pt=11, bold=r.bold)


# ---------------------------------------------------------------------------
# Step 3: open, restyle, save
# ---------------------------------------------------------------------------
print('[2/3] Restyling document per dissertation guidelines...')

doc = Document(DOCX_FILE)

# ---- Find cover-page boundary: paragraphs before the first "Declaration" heading
declaration_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower() == 'declaration':
        declaration_idx = i
        break
if declaration_idx is None:
    declaration_idx = 13  # safety fallback
print(f'      cover page ends at paragraph {declaration_idx}')

# ---- Consolidate sections (remove every intermediate sectPr) ----
body = doc.element.body
for p_el in body.iter(W('w:p')):
    pPr = p_el.find(W('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(W('w:sectPr'))
        if sectPr is not None:
            pPr.remove(sectPr)

# Now there is exactly one section (the trailing one in body) - set A4 + 1" margins
section = doc.sections[0]
section.page_width = Inches(8.27)     # A4
section.page_height = Inches(11.69)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

# ---- Walk all paragraphs and classify ----
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()

    # Cover page: keep centered, TNR
    if i <= declaration_idx - 1:
        if not text:
            continue
        style_cover(p)
        continue

    if not text:
        # blank paragraph: minimal cleanup, do not let it contribute extra space
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        for r in p.runs:
            set_run_font(r, size_pt=12)
        continue

    max_size = get_max_size(p)
    bold = all_runs_bold(p)

    # Captions: anything that starts with "Figure N" or "Table N"
    if is_caption(text):
        style_caption(p)
        continue

    # TOC / List of Figures / List of Tables entries: keep left-aligned, single line
    if is_toc_entry(text):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        for r in p.runs:
            set_run_font(r, size_pt=12, bold=r.bold)
        continue

    # Explicit front-matter / appendix chapter headings
    if text.lower() in KNOWN_MAIN_HEADINGS:
        style_main_heading(p)
        continue

    # Main heading: source font >= 16 pt and bold (chapter / part titles)
    if max_size is not None and max_size >= 16 and bold:
        style_main_heading(p)
        continue

    # Sub-heading: source font 13-16 pt, bold
    if max_size is not None and 13 <= max_size < 16 and bold:
        style_sub_heading(p)
        continue

    # Title-like centered lines after cover (rare): keep centered, 14pt
    if max_size is not None and max_size >= 13.9:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        for r in p.runs:
            set_run_font(r, size_pt=14, bold=r.bold)
        continue

    # Default: body
    style_body(p)

# ---- Style every paragraph inside every table cell ----
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                style_table_cell_paragraph(p)

# ---- Set document default font to TNR via styles.xml ----
styles_element = doc.styles.element
rPrDefault = styles_element.find(qn('w:docDefaults') + '/' + qn('w:rPrDefault'))
if rPrDefault is not None:
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)
    # font
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), 'Times New Roman')
    # size 24 half-pt = 12 pt
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '24')

# Style "Normal" -> TNR 12pt as well (some Word builds key off this)
try:
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), 'Times New Roman')
except Exception as e:
    print('  (styles[Normal] tweak skipped:', e, ')')

doc.save(DOCX_FILE)

# ---- Quick stats ----
print('[3/3] Done.')
doc2 = Document(DOCX_FILE)
print(f'      Sections: {len(doc2.sections)}')
print(f'      Paragraphs: {len(doc2.paragraphs)}')
print(f'      Tables: {len(doc2.tables)}')
print(f'      Inline images: {len(doc2.inline_shapes)}')
import os
print(f'      File size: {os.path.getsize(DOCX_FILE)/1024:.1f} KB')
