from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_paragraph_spacing(paragraph, space_after_pt=12, line_spacing_rule=None, line_spacing_val=None):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(0)
    if line_spacing_val:
        from docx.shared import Pt as Pt2
        from docx.oxml.ns import qn as qn2
        from docx.oxml import OxmlElement as OXE
        # 1.5 line spacing = 276 twips (240 * 1.5 = 360 but Word uses 240 per line)
        pPr = paragraph._p.get_or_add_pPr()
        lSpacing = OxmlElement('w:spacing')
        lSpacing.set(qn('w:line'), str(line_spacing_val))
        lSpacing.set(qn('w:lineRule'), 'auto')
        # Remove existing spacing element if any
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(lSpacing)

def create_acknowledgement_doc(output_path):
    doc = Document()

    # ----- PAGE SETUP: A4, 1-inch margins -----
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    margin = Cm(2.54)
    section.top_margin    = margin
    section.bottom_margin = margin
    section.left_margin   = margin
    section.right_margin  = margin

    # ----- Default document font -----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Fix for East-Asian font fallback in python-docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),   'Times New Roman')
    rFonts.set(qn('w:cs'),      'Times New Roman')

    # ----- HEADING: ACKNOWLEDGEMENT -----
    heading_para = doc.add_paragraph()
    heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_para.paragraph_format.space_before = Pt(0)
    heading_para.paragraph_format.space_after  = Pt(12)

    # 1.5 line spacing for heading too
    pPr = heading_para._p.get_or_add_pPr()
    lSpacing = OxmlElement('w:spacing')
    lSpacing.set(qn('w:line'), '360')   # 240 * 1.5
    lSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(lSpacing)

    run = heading_para.add_run('ACKNOWLEDGEMENT')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    # Ensure font applied via XML
    rPr2 = run._r.get_or_add_rPr()
    rFonts2 = rPr2.find(qn('w:rFonts'))
    if rFonts2 is None:
        rFonts2 = OxmlElement('w:rFonts')
        rPr2.insert(0, rFonts2)
    rFonts2.set(qn('w:ascii'), 'Times New Roman')
    rFonts2.set(qn('w:hAnsi'), 'Times New Roman')

    # ----- BODY PARAGRAPHS -----
    body_paragraphs = [
        "In the name of Allah, the Most Gracious, the Most Merciful.",

        "First and foremost, the author expresses profound gratitude to Allah (SWT) for His infinite mercy, guidance, blessings, and strength throughout this academic journey. By His grace, the author was granted the perseverance, knowledge, and opportunities necessary to successfully complete this project.",

        "The author extends heartfelt appreciation to Mr. Rashid Akber and Mrs. Kaniz Sakeena Rashid for their unwavering love, endless sacrifices, constant encouragement, prayers, and unconditional support. Their guidance, patience, and steadfast belief have been the foundation of the author's personal and academic growth and a continuous source of motivation throughout this endeavor.",

        "Special gratitude is extended to Ms. Rida Fatima for her immense patience, understanding, encouragement, and unwavering moral support. Her constant reassurance, kindness, and belief in the author provided strength and motivation during challenging phases of this project.",

        "The author is deeply indebted to Prof. Mohammad Muzammil for his invaluable guidance, insightful suggestions, constructive feedback, and continuous encouragement throughout the course of this work. His mentorship, expertise, and academic support were instrumental in shaping the direction and successful completion of this project.",

        "Sincere appreciation is also extended to Mr. Sheikh Noorul, project partner, for his dedication, cooperation, and collaborative efforts throughout the development of this project. His teamwork, valuable contributions, and willingness to overcome challenges together greatly facilitated the achievement of the project objectives.",

        "The author would also like to express sincere gratitude to all faculty members of the Department of Mechanical Engineering for their guidance, encouragement, and support throughout the academic program. Their knowledge and assistance have contributed significantly to the author's learning and development.",

        "Finally, heartfelt thanks are extended to family members, friends, and well-wishers whose encouragement, understanding, prayers, and support provided constant motivation and inspiration throughout this journey. Their confidence and goodwill contributed greatly to the successful completion of this work.",
    ]

    for text in body_paragraphs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(12)

        # 1.5 line spacing
        pPr = para._p.get_or_add_pPr()
        lSpacing = OxmlElement('w:spacing')
        lSpacing.set(qn('w:line'), '360')   # 240 * 1.5
        lSpacing.set(qn('w:lineRule'), 'auto')
        pPr.append(lSpacing)

        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        rPr3 = run._r.get_or_add_rPr()
        rFonts3 = rPr3.find(qn('w:rFonts'))
        if rFonts3 is None:
            rFonts3 = OxmlElement('w:rFonts')
            rPr3.insert(0, rFonts3)
        rFonts3.set(qn('w:ascii'), 'Times New Roman')
        rFonts3.set(qn('w:hAnsi'), 'Times New Roman')

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == '__main__':
    create_acknowledgement_doc('/projects/sandbox/12/Acknowledgement.docx')
